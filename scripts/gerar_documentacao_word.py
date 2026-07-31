from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "DOCUMENTACAO_COMPLETA.md"
OUTPUT = ROOT / "docs" / "DOCUMENTACAO_COMPLETA_RADAR_VIAGEM_SEGURA.docx"

NAVY = "0B1F33"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
CYAN = "38BDF8"
TEXT = "243247"
MUTED = "64748B"
LIGHT = "F5F8FC"
WHITE = "FFFFFF"
BORDER = "C9D6E5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=130, bottom=120, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Clique com o botão direito e selecione Atualizar Campo para atualizar o sumário."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^]]+\]\([^)]+\))")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Cascadia Mono"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BLUE)
        else:
            label, url = re.match(r"\[([^]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.underline = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "101827")
    set_cell_margins(cell, 180, 220, 180, 220)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Cascadia Mono"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("D9E6F5")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    data = [rows[0]] + rows[2:]
    table = document.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, values in enumerate(data):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_shading(cell, NAVY if row_index == 0 else (LIGHT if row_index % 2 == 0 else WHITE))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value.strip().replace("**", ""))
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Heading 1", 19, NAVY, 18, 8),
        ("Heading 2", 14, BLUE, 14, 6),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(10.5)
        styles[name].font.color.rgb = RGBColor.from_string(TEXT)

    if "Destaque" not in styles:
        highlight = styles.add_style("Destaque", WD_STYLE_TYPE.PARAGRAPH)
        highlight.font.name = "Aptos"
        highlight.font.size = Pt(10)
        highlight.font.italic = True
        highlight.font.color.rgb = RGBColor.from_string(NAVY)
        highlight.paragraph_format.left_indent = Cm(0.6)
        highlight.paragraph_format.right_indent = Cm(0.6)
        highlight.paragraph_format.space_before = Pt(6)
        highlight.paragraph_format.space_after = Pt(10)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    header = section.header.paragraphs[0]
    header.text = "RADAR DE VIAGEM SEGURA   |   DOCUMENTAÇÃO DO PROJETO"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
    add_page_number(section.footer.paragraphs[0])


def add_cover(document: Document) -> None:
    accent = document.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = accent.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell, 50, 50, 50, 50)
    cell.paragraphs[0].add_run(" ")

    document.add_paragraph().paragraph_format.space_after = Pt(54)
    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("PROJETO ACADÊMICO  •  VERSÃO 1.0")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("RADAR DE\nVIAGEM SEGURA")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(4)
    subtitle.paragraph_format.space_after = Pt(32)
    run = subtitle.add_run("Documentação completa do sistema")
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    callout.columns[0].width = Cm(14.5)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, 260, 340, 260, 340)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Análise histórica de acidentes, apoio ao planejamento de viagens e exploração visual de dados viários.")
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    document.add_paragraph().paragraph_format.space_after = Pt(58)
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metadata.add_run("Python  •  Streamlit  •  Pandas  •  Plotly  •  scikit-learn\n31 de julho de 2026")
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_page_break()


def render_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1]):
            table_rows = []
            while index < len(lines) and lines[index].startswith("|"):
                table_rows.append([value.strip() for value in lines[index].strip("|").split("|")])
                index += 1
            add_markdown_table(document, table_rows)
            continue
        if not line.strip() or line.strip() == "---":
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1:
                index += 1
                continue
            paragraph = document.add_heading(text, level=level - 1)
            paragraph.paragraph_format.page_break_before = level == 2
            index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Destaque")
            add_inline(paragraph, line[2:])
            index += 1
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet or numbered:
            paragraph = document.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue
        paragraph = document.add_paragraph()
        add_inline(paragraph, line.rstrip("  "))
        index += 1


def main() -> None:
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)

    document.add_heading("Sumário", level=1)
    toc = document.add_paragraph()
    add_toc(toc)
    document.add_page_break()

    source = SOURCE.read_text(encoding="utf-8")
    render_markdown(document, source)

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    properties = document.core_properties
    properties.title = "Radar de Viagem Segura — Documentação completa"
    properties.subject = "Documentação técnica e manual de uso"
    properties.author = "Projeto Radar de Viagem Segura"
    properties.keywords = "TCC, segurança viária, Streamlit, acidentes, Python"

    document.save(OUTPUT)
    print(f"Documento gerado: {OUTPUT}")


if __name__ == "__main__":
    main()
